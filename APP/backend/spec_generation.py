import os
import re
import json
import uuid
import threading
import logging
from typing import Dict, Any, Tuple, List, Optional
from datetime import datetime
from llm_service import LLMService
from vector_store import VectorStore

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.normpath(os.path.join(BASE_DIR, "../.."))

SPEC_DIR = os.path.join(PROJECT_ROOT, "design/spec")

_LLM_CONFIG_PATH = os.path.join(BASE_DIR, 'llm_config.json')


def _load_file_llm_config() -> dict:
    """从 llm_config.json 读取当前生效的 LLM 配置"""
    try:
        with open(_LLM_CONFIG_PATH, encoding='utf-8') as f:
            full = json.load(f)
        provider = full.get('last_provider', '')
        if not provider or provider == 'none':
            return {}
        pcfg = full.get(provider, {})
        return {
            'provider': provider,
            'api_key': pcfg.get('api_key', ''),
            'model_name': pcfg.get('model_name', ''),
            'base_url': pcfg.get('base_url', ''),
        }
    except Exception:
        return {}
TEMPLATE_DIR = os.path.join(PROJECT_ROOT, "design/template")
EXPORT_DIR = os.path.join(PROJECT_ROOT, "conclusion/exports")

# Ensure directories exist
for d in [SPEC_DIR, TEMPLATE_DIR, EXPORT_DIR]:
    os.makedirs(d, exist_ok=True)


class SpecGenerator:
    """需求规格文档生成器 - 支持Markdown、DOCX、PDF格式"""

    def __init__(self, vector_store: VectorStore, llm_service: LLMService):
        self.vector_store = vector_store
        self.llm_service = llm_service
        self.tasks: Dict[str, Dict] = {}  # 异步任务状态
        self.logger = logging.getLogger("ai_ticket.spec_generation")

    def _load_template(self, filename: str) -> str:
        filepath = os.path.join(TEMPLATE_DIR, filename)
        if not os.path.exists(filepath):
            return ""
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()

    def _clean_filename(self, title: str) -> str:
        """清理标题，生成安全的文件名"""
        # 替换不安全字符
        safe = title.replace('/', '_').replace('\\', '_').replace(':', '_')
        safe = re.sub(r'[<>"|?*]', '', safe)
        safe = re.sub(r'\s+', '_', safe)
        return safe[:50]  # 限制长度

    def _build_section_plan(self, doc_type: str) -> List[Dict[str, Any]]:
        if doc_type == "summary":
            return [
                {"id": "background_and_goal", "title": "需求背景与目标", "required": True, "required_keywords": ["背景", "目标"]},
                {"id": "business_scenarios", "title": "业务场景", "required": True, "required_keywords": ["场景"]},
                {"id": "business_flow", "title": "业务流程", "required": True, "required_keywords": ["流程"]},
                {"id": "product_solution", "title": "产品方案", "required": True, "required_keywords": ["方案"]},
                {"id": "risks_and_boundaries", "title": "边界与风险", "required": False, "required_keywords": ["风险", "边界"]},
            ]
        if doc_type == "detail":
            return [
                {"id": "feature_overview", "title": "功能概述", "required": True, "required_keywords": ["功能"]},
                {"id": "interaction_flow", "title": "交互流程", "required": True, "required_keywords": ["交互", "流程"]},
                {"id": "page_and_field_details", "title": "页面与字段说明", "required": True, "required_keywords": ["字段", "页面"]},
                {"id": "api_and_data_model", "title": "接口与数据结构", "required": True, "required_keywords": ["接口", "数据"]},
                {"id": "compatibility_and_migration", "title": "兼容性与迁移", "required": False, "required_keywords": ["兼容", "迁移"]},
                {"id": "acceptance_and_test_cases", "title": "验收标准与测试要点", "required": True, "required_keywords": ["验收", "测试"]},
            ]
        raise ValueError(f"Unsupported doc_type: {doc_type}")

    def _create_generation_task(self, task_id: str, req_id: str) -> Dict[str, Any]:
        task = {
            "task_id": task_id,
            "req_id": req_id,
            "status": "pending",
            "is_executing": False,
            "documents": {},
            "error_summary": "",
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
        }
        for doc_type in ["summary", "detail"]:
            sections = []
            for section in self._build_section_plan(doc_type):
                section_state = dict(section)
                section_state.update({
                    "status": "pending",
                    "content": "",
                    "error": "",
                    "attempts": 0,
                })
                sections.append(section_state)
            task["documents"][doc_type] = {
                "doc_type": doc_type,
                "status": "pending",
                "is_executing": False,
                "sections": sections,
                "output_file": None,
            }
        return task

    def _validate_section_output(self, section: Dict[str, Any], content: str) -> Dict[str, Any]:
        text = (content or "").strip()
        lowered = text.lower()
        invalid_markers = [
            "模型调用失败",
            "error code:",
            "invalid_authentication_error",
            "api key",
            "hello!",
            "how can i help you today",
        ]

        if not text:
            return {"valid": False, "reason": "empty_output"}
        if any(marker in lowered for marker in [m.lower() for m in invalid_markers]):
            return {"valid": False, "reason": "bad_output"}

        required_keywords = section.get("required_keywords") or []
        if required_keywords and not any(keyword in text for keyword in required_keywords):
            return {"valid": False, "reason": "missing_keywords"}

        return {"valid": True, "reason": ""}

    def _summarize_document_status(self, document: Dict[str, Any]) -> Dict[str, Any]:
        sections = document.get("sections", [])
        completed = sum(1 for section in sections if section.get("status") == "completed")
        failed = sum(1 for section in sections if section.get("status") == "failed")
        running = sum(1 for section in sections if section.get("status") == "running")
        pending = sum(1 for section in sections if section.get("status") == "pending")
        required_failed = any(
            section.get("required", True) and section.get("status") == "failed"
            for section in sections
        )

        if document.get("is_executing") and (running > 0 or pending > 0):
            status = "running"
        elif required_failed:
            status = "failed"
        elif failed:
            status = "partial"
        elif pending:
            status = "pending"
        else:
            status = "completed"

        return {
            "status": status,
            "completed_sections": completed,
            "failed_sections": failed,
            "running_sections": running,
            "pending_sections": pending,
        }

    def _summarize_task_status(self, task: Dict[str, Any]) -> Dict[str, Any]:
        documents = {}
        has_failed = False
        has_partial = False
        all_completed = True
        any_running = False

        for doc_type, document in task.get("documents", {}).items():
            summary = self._summarize_document_status(document)
            document["status"] = summary["status"]
            documents[doc_type] = summary
            has_failed = has_failed or summary["status"] == "failed"
            has_partial = has_partial or summary["status"] == "partial"
            all_completed = all_completed and summary["status"] == "completed"
            any_running = any_running or summary["status"] == "running"

        if any_running or task.get("is_executing"):
            status = "running"
        elif has_failed or has_partial:
            status = "partial"
        elif all_completed:
            status = "completed"
        else:
            status = "pending"

        task["status"] = status
        task["updated_at"] = datetime.now().isoformat()
        return {"status": status, "documents": documents}

    def _get_retryable_section_ids(self, document: Dict[str, Any]) -> List[str]:
        return [
            section["id"]
            for section in document.get("sections", [])
            if section.get("status") == "failed"
        ]

    def _clean_output(self, txt: str) -> str:
        txt = (txt or "").strip()
        if txt.startswith("```markdown"):
            txt = txt[11:]
        elif txt.startswith("```"):
            txt = txt[3:]
        if txt.endswith("```"):
            txt = txt[:-3]
        txt = re.sub(r"<think>.*?</think>\s*", "", txt, flags=re.DOTALL | re.IGNORECASE)
        return txt.strip()

    def _assemble_document_markdown(self, document_title: str, sections: List[Dict[str, Any]]) -> str:
        parts = [f"# {document_title}"]
        for section in sections:
            content = (section.get("content") or "").strip()
            if not content.startswith("## "):
                content = f"## {section['title']}\n{content}"
            parts.append(content)
        return "\n\n".join(parts).strip() + "\n"

    def _build_generation_context(self, req_id: str, req: Dict[str, Any], version: str, records: List[Dict[str, Any]]) -> str:
        ai_analysis = req.get('ai_analysis', {})
        evidence_bundle = ai_analysis.get('evidence_bundle', {}) or {}
        evidence_lines = []
        for item in (evidence_bundle.get('evidence') or [])[:5]:
            source_kind = item.get('source_kind', 'unknown')
            name = item.get('name', '')
            summary = item.get('summary', '')
            evidence_lines.append(f"- [{source_kind}] {name}: {summary}")

        topic_names = ai_analysis.get('topic_names') or evidence_bundle.get('topic_names') or []
        suggested_sections = evidence_bundle.get('suggested_sections') or []
        open_questions = evidence_bundle.get('open_questions') or []

        return f"""你是一名资深产品经理和系统架构师。
请你根据以下原始需求和前期AI的根因分析、模块定位，结合质量评审意见，生成规范的产品需求文档。

【需求基础信息】
编号: {req_id}
标题: {req.get('title', '')}
原工单描述: {req.get('description', '')[:2000]}
计划落地版本: {version}

【前期智能分析结果】
根因分析: {req.get('ai_analysis', {}).get('root_cause', '')}
影响模块: {req.get('ai_analysis', {}).get('module', '')}
MVP建议: {req.get('ai_analysis', {}).get('mvp_suggestion', '')}

【证据包】
主题名称: {", ".join(topic_names) if topic_names else "无"}
建议章节: {", ".join(suggested_sections) if suggested_sections else "无"}
待确认项: {"；".join(open_questions) if open_questions else "无"}
证据摘要:
{os.linesep.join(evidence_lines) if evidence_lines else "- 当前未命中外部证据，请保守表达并明确待确认项。"}

【评审意见】
{records[-1].get('comments', '') if records else '无'}
"""

    def generate_draft_artifact(self, req_id: str, draft_type: str) -> Dict[str, Any]:
        if draft_type not in {"summary", "detail"}:
            raise ValueError("draft_type must be 'summary' or 'detail'")

        inputs = self._resolve_generation_inputs(req_id)
        document_name = "概要需求初稿" if draft_type == "summary" else "详细需求初稿"
        template = inputs["summary_template"] if draft_type == "summary" else inputs["detail_template"]
        sections = self._build_section_plan(draft_type)

        success, message, markdown = self._generate_document_content(
            document_name=document_name,
            sections=sections,
            base_context=inputs["base_context"],
            template=template,
            api_key=inputs["api_key"],
            provider=inputs["provider"],
            model_name=inputs["model_name"],
            base_url=inputs["base_url"],
        )
        if not success:
            raise ValueError(message)

        suffix = "概要需求初稿" if draft_type == "summary" else "详细需求初稿"
        filename = f"{inputs['base_filename']}-{suffix}.md"
        spec_path = os.path.join(SPEC_DIR, filename)
        with open(spec_path, "w", encoding="utf-8") as handle:
            handle.write(markdown)

        ai_analysis = inputs["req"].get("ai_analysis", {}) or {}
        return {
            "req_id": req_id,
            "draft_type": draft_type,
            "spec_file": filename,
            "spec_path": spec_path,
            "source": "req_pool",
            "title": inputs["req"].get("title", ""),
            "module_hint": ai_analysis.get("module_hint") or ai_analysis.get("module", ""),
            "topic_names": ai_analysis.get("topic_names", []),
            "evidence_bundle": ai_analysis.get("evidence_bundle", {}),
            "created_at": datetime.now().isoformat(),
        }

    def _build_section_prompt(self, base_context: str, document_name: str, section: Dict[str, Any], template: str) -> str:
        return f"""{base_context}

你现在只需要编写《{document_name}》中的一个章节，不要输出整份文档。

【当前章节】
标题: {section['title']}

【章节要求】
- 只输出当前章节的 Markdown 内容
- 优先围绕标题组织内容
- 内容必须与当前需求直接相关
- 不要输出问候语、说明文字或与章节无关的废话

【模板参考】
{template}
"""

    def _generate_document_content(
        self,
        document_name: str,
        sections: List[Dict[str, Any]],
        base_context: str,
        template: str,
        api_key: str,
        provider: str,
        model_name: str,
        base_url: str = "",
    ) -> Tuple[bool, str, str]:
        generated_sections = []

        for section in sections:
            prompt = self._build_section_prompt(base_context, document_name, section, template)
            content = self.llm_service.call_llm(
                prompt=prompt,
                api_key=api_key,
                provider=provider,
                model_name=model_name,
                base_url=base_url if base_url else None
            )
            cleaned = self._clean_output(content)
            validation = self._validate_section_output(section, cleaned)
            if not validation["valid"]:
                return False, f"Invalid output for {document_name}/{section['title']}: {validation['reason']}", ""
            generated_sections.append({
                "title": section["title"],
                "content": cleaned,
            })

        return True, "", self._assemble_document_markdown(document_name, generated_sections)

    def _resolve_generation_inputs(self, req_id: str) -> Dict[str, Any]:
        req = self.vector_store.get_requirement(req_id)
        if not req:
            raise ValueError("Requirement not found.")

        if req.get('status') != 'scheduled':
            raise ValueError("Requirement must be in 'scheduled' status.")

        version = "V_Next"
        records = req.get('review_records') or []
        for r in reversed(records):
            if r.get('expected_version'):
                version = r.get('expected_version')
                break

        summary_tpl = self._load_template("概要需求模板.md")
        detail_tpl = self._load_template("详细需求模板.md")
        if not summary_tpl or not detail_tpl:
            raise ValueError("Templates '概要需求模板.md' or '详细需求模板.md' not found.")

        api_key = os.environ.get("LLM_API_KEY", "")
        provider = os.environ.get("LLM_PROVIDER", "")
        model_name = os.environ.get("LLM_MODEL_NAME", "")
        base_url = os.environ.get("LLM_BASE_URL", "")

        if hasattr(self.llm_service, 'default_api_key'):
            api_key = self.llm_service.default_api_key or api_key
        if hasattr(self.llm_service, 'default_provider'):
            provider = self.llm_service.default_provider or provider
        if hasattr(self.llm_service, 'default_model_name'):
            model_name = self.llm_service.default_model_name or model_name
        if hasattr(self.llm_service, 'default_base_url'):
            base_url = self.llm_service.default_base_url or base_url

        # 最终降级：读取 llm_config.json
        if not api_key:
            file_cfg = _load_file_llm_config()
            api_key = file_cfg.get('api_key', '')
            provider = provider or file_cfg.get('provider', 'gemini')
            model_name = model_name or file_cfg.get('model_name', '')
            base_url = base_url or file_cfg.get('base_url', '')

        if not model_name:
            if provider == "gemini":
                model_name = "gemini-2.5-pro"
            elif provider == "openai":
                model_name = "gpt-4"
            elif provider == "deepseek":
                model_name = "deepseek-chat"

        if not api_key:
            raise ValueError("LLM API key not configured.")

        safe_title = self._clean_filename(req.get('title', 'Unknown'))
        base_context = self._build_generation_context(req_id, req, version, records)
        base_filename = f"{version}-{req_id}-{safe_title}"

        return {
            "req": req,
            "version": version,
            "records": records,
            "summary_template": summary_tpl,
            "detail_template": detail_tpl,
            "api_key": api_key,
            "provider": provider,
            "model_name": model_name,
            "base_url": base_url,
            "safe_title": safe_title,
            "base_context": base_context,
            "base_filename": base_filename,
        }

    def _mark_task_failed(self, task: Dict[str, Any], message: str):
        for document in task["documents"].values():
            for section in document["sections"]:
                if section.get("status") != "completed":
                    section["status"] = "failed"
                    section["error"] = message
        task["error_summary"] = message
        self._summarize_task_status(task)

    def _generate_section_for_task(
        self,
        task: Dict[str, Any],
        doc_type: str,
        document_name: str,
        section: Dict[str, Any],
        base_context: str,
        template: str,
        api_key: str,
        provider: str,
        model_name: str,
        base_url: str = "",
    ):
        document = task["documents"][doc_type]
        document["status"] = "running"
        section["status"] = "running"
        section["error"] = ""
        task["updated_at"] = datetime.now().isoformat()
        self._summarize_task_status(task)

        try:
            prompt = self._build_section_prompt(base_context, document_name, section, template)
            content = self.llm_service.call_llm(
                prompt=prompt,
                api_key=api_key,
                provider=provider,
                model_name=model_name,
                base_url=base_url if base_url else None
            )
            section["attempts"] = section.get("attempts", 0) + 1
            cleaned = self._clean_output(content)
            validation = self._validate_section_output(section, cleaned)
            if validation["valid"]:
                section["status"] = "completed"
                section["content"] = cleaned
                section["error"] = ""
            else:
                section["status"] = "failed"
                section["content"] = ""
                section["error"] = validation["reason"]
        except Exception as e:
            section["attempts"] = section.get("attempts", 0) + 1
            section["status"] = "failed"
            section["content"] = ""
            section["error"] = str(e)
        finally:
            task["updated_at"] = datetime.now().isoformat()
            self._summarize_task_status(task)

    def _run_document_sections(
        self,
        task: Dict[str, Any],
        doc_type: str,
        template: str,
        base_context: str,
        api_key: str,
        provider: str,
        model_name: str,
        base_url: str = "",
        section_ids: Optional[List[str]] = None,
    ):
        document_name = "概要需求" if doc_type == "summary" else "详细需求"
        document = task["documents"][doc_type]
        document["status"] = "running"
        document["is_executing"] = True

        try:
            for section in document["sections"]:
                if section_ids is not None and section["id"] not in section_ids:
                    continue
                self._generate_section_for_task(
                    task=task,
                    doc_type=doc_type,
                    document_name=document_name,
                    section=section,
                    base_context=base_context,
                    template=template,
                    api_key=api_key,
                    provider=provider,
                    model_name=model_name,
                    base_url=base_url,
                )
        finally:
            document["is_executing"] = False
            self._summarize_task_status(task)

    def _attach_output_files(self, task: Dict[str, Any], output_files: List[str]):
        for doc_type, document in task["documents"].items():
            suffix = "概要需求" if doc_type == "summary" else "详细需求"
            doc_files = [name for name in output_files if suffix in name]
            document["output_files"] = doc_files
            preferred = next((name for name in doc_files if name.endswith(".md")), None)
            document["output_file"] = preferred or (doc_files[0] if doc_files else None)

    def _write_output_files(
        self,
        summary_content: str,
        detail_content: str,
        formats: List[str],
        base_filename: str,
        safe_title: str,
        version: str,
    ) -> Dict[str, Any]:
        output_files = []

        for fmt in formats:
            if fmt == "md":
                summary_md = f"{base_filename}-概要需求.md"
                detail_md = f"{base_filename}-详细需求.md"

                with open(os.path.join(SPEC_DIR, summary_md), 'w', encoding='utf-8') as f:
                    f.write(summary_content)
                with open(os.path.join(SPEC_DIR, detail_md), 'w', encoding='utf-8') as f:
                    f.write(detail_content)

                output_files.append(summary_md)
                output_files.append(detail_md)

            elif fmt == "docx":
                summary_docx = os.path.join(SPEC_DIR, f"{base_filename}-概要需求.docx")
                detail_docx = os.path.join(SPEC_DIR, f"{base_filename}-详细需求.docx")

                if self._generate_docx(summary_content, summary_docx, f"{safe_title} - 概要需求"):
                    output_files.append(f"{base_filename}-概要需求.docx")
                if self._generate_docx(detail_content, detail_docx, f"{safe_title} - 详细需求"):
                    output_files.append(f"{base_filename}-详细需求.docx")

            elif fmt == "pdf":
                summary_pdf = os.path.join(SPEC_DIR, f"{base_filename}-概要需求.pdf")
                detail_pdf = os.path.join(SPEC_DIR, f"{base_filename}-详细需求.pdf")

                if self._generate_pdf(summary_content, summary_pdf, f"{safe_title} - 概要需求"):
                    output_files.append(f"{base_filename}-概要需求.pdf")
                if self._generate_pdf(detail_content, detail_pdf, f"{safe_title} - 详细需求"):
                    output_files.append(f"{base_filename}-详细需求.pdf")

        return {
            "base_filename": base_filename,
            "output_files": output_files,
            "version": version,
            "spec_dir": SPEC_DIR
        }

    def _finalize_generation_task(
        self,
        task: Dict[str, Any],
        inputs: Dict[str, Any],
        formats: List[str],
    ):
        if task.get("status") != "completed":
            return

        summary_sections = [
            {"title": section["title"], "content": section["content"]}
            for section in task["documents"]["summary"]["sections"]
            if section.get("status") == "completed"
        ]
        detail_sections = [
            {"title": section["title"], "content": section["content"]}
            for section in task["documents"]["detail"]["sections"]
            if section.get("status") == "completed"
        ]

        summary_content = self._assemble_document_markdown("概要需求", summary_sections)
        detail_content = self._assemble_document_markdown("详细需求", detail_sections)

        task["documents"]["summary"]["markdown"] = summary_content
        task["documents"]["detail"]["markdown"] = detail_content

        output_info = self._write_output_files(
            summary_content=summary_content,
            detail_content=detail_content,
            formats=formats,
            base_filename=inputs["base_filename"],
            safe_title=inputs["safe_title"],
            version=inputs["version"],
        )
        task["artifacts"] = output_info
        self._attach_output_files(task, output_info["output_files"])
        task["error_summary"] = ""

    def start_generation_task(self, req_id: str, formats: List[str] = None) -> str:
        if formats is None:
            formats = ["md"]

        task_id = str(uuid.uuid4())
        task = self._create_generation_task(task_id, req_id)
        task["requested_formats"] = formats
        self.tasks[task_id] = task

        worker = threading.Thread(
            target=self._run_generation_task,
            args=(task_id, req_id, formats),
            daemon=True,
        )
        self.logger.info("start_generation_task before worker.start req_id=%s task_id=%s formats=%s", req_id, task_id, formats)
        worker.start()
        self.logger.info("start_generation_task after worker.start req_id=%s task_id=%s", req_id, task_id)
        return task_id

    def _run_generation_task(self, task_id: str, req_id: str, formats: List[str]):
        task = self.tasks.get(task_id)
        if not task:
            return

        self.logger.info("_run_generation_task entered req_id=%s task_id=%s", req_id, task_id)
        task["status"] = "running"
        task["is_executing"] = True
        task["updated_at"] = datetime.now().isoformat()

        try:
            inputs = self._resolve_generation_inputs(req_id)
            self._run_document_sections(
                task=task,
                doc_type="summary",
                template=inputs["summary_template"],
                base_context=inputs["base_context"],
                api_key=inputs["api_key"],
                provider=inputs["provider"],
                model_name=inputs["model_name"],
                base_url=inputs["base_url"],
            )
            self._run_document_sections(
                task=task,
                doc_type="detail",
                template=inputs["detail_template"],
                base_context=inputs["base_context"],
                api_key=inputs["api_key"],
                provider=inputs["provider"],
                model_name=inputs["model_name"],
                base_url=inputs["base_url"],
            )
            task["is_executing"] = False
            self._summarize_task_status(task)
            if task["status"] == "completed":
                self._finalize_generation_task(task, inputs, formats)
            else:
                errors = [
                    f"{doc_type}:{section['id']}:{section['error']}"
                    for doc_type, document in task["documents"].items()
                    for section in document["sections"]
                    if section.get("status") in {"failed", "partial"} and section.get("error")
                ]
                task["error_summary"] = "; ".join(errors)
                task["artifacts"] = {
                    "base_filename": inputs["base_filename"],
                    "output_files": [],
                    "version": inputs["version"],
                    "spec_dir": SPEC_DIR,
                }
        except ValueError as e:
            self._mark_task_failed(task, str(e))
        except Exception as e:
            self._mark_task_failed(task, str(e))
        finally:
            task["is_executing"] = False
            for document in task["documents"].values():
                document["is_executing"] = False
            self._summarize_task_status(task)

    def retry_generation_task(self, task_id: str) -> Optional[Dict[str, Any]]:
        task = self.tasks.get(task_id)
        if not task:
            return None

        section_ids_by_doc = {}
        for doc_type, document in task["documents"].items():
            retryable_ids = self._get_retryable_section_ids(document)
            if retryable_ids:
                section_ids_by_doc[doc_type] = retryable_ids

        if not section_ids_by_doc:
            return {
                "task_id": task_id,
                "retried_sections": {},
            }

        worker = threading.Thread(
            target=self._retry_generation_task,
            args=(task_id, section_ids_by_doc),
            daemon=True,
        )
        worker.start()
        return {
            "task_id": task_id,
            "retried_sections": section_ids_by_doc,
        }

    def _retry_generation_task(self, task_id: str, section_ids_by_doc: Dict[str, List[str]]):
        task = self.tasks.get(task_id)
        if not task:
            return

        try:
            task["is_executing"] = True
            inputs = self._resolve_generation_inputs(task["req_id"])

            for doc_type, section_ids in section_ids_by_doc.items():
                document = task["documents"].get(doc_type)
                if not document:
                    continue
                for section in document["sections"]:
                    if section["id"] in section_ids:
                        section["status"] = "pending"
                        section["error"] = ""
                self._run_document_sections(
                    task=task,
                    doc_type=doc_type,
                    template=inputs["summary_template"] if doc_type == "summary" else inputs["detail_template"],
                    base_context=inputs["base_context"],
                    api_key=inputs["api_key"],
                    provider=inputs["provider"],
                    model_name=inputs["model_name"],
                    base_url=inputs["base_url"],
                    section_ids=section_ids,
                )

            task["is_executing"] = False
            self._summarize_task_status(task)
            if task["status"] == "completed":
                self._finalize_generation_task(task, inputs, task.get("requested_formats") or ["md"])
        except ValueError as e:
            self._mark_task_failed(task, str(e))
        except Exception as e:
            self._mark_task_failed(task, str(e))
        finally:
            task["is_executing"] = False
            for document in task["documents"].values():
                document["is_executing"] = False
            self._summarize_task_status(task)

    def _generate_docx(self, content: str, filepath: str, title: str):
        """从Markdown内容生成DOCX文档"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Add title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Parse markdown content
            lines = content.split('\n')
            current_table = []
            in_code_block = False
            code_content = []
            code_lang = ""

            for line in lines:
                # Handle code blocks
                if line.strip().startswith('```'):
                    if in_code_block:
                        # End code block
                        if code_content:
                            para = doc.add_paragraph()
                            run = para.add_run('\n'.join(code_content))
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                        code_content = []
                        in_code_block = False
                        code_lang = ""
                    else:
                        # Start code block
                        in_code_block = True
                        code_lang = line.strip()[3:].strip()
                    continue

                if in_code_block:
                    code_content.append(line)
                    continue

                # Handle headers
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                # Handle table rows
                elif line.startswith('|') and '|' in line[1:]:
                    current_table.append(line)
                else:
                    # Flush table if exists
                    if current_table:
                        self._add_table_to_docx(doc, current_table)
                        current_table = []

                    # Handle list items
                    if line.strip().startswith('- '):
                        doc.add_paragraph(line.strip()[2:], style='List Bullet')
                    elif re.match(r'^\d+\. ', line.strip()):
                        doc.add_paragraph(re.sub(r'^\d+\. ', '', line.strip()), style='List Number')
                    elif line.strip():
                        doc.add_paragraph(line.strip())

            # Flush remaining table
            if current_table:
                self._add_table_to_docx(doc, current_table)

            doc.save(filepath)
            return True
        except Exception as e:
            print(f"[SpecGenerator] Error generating DOCX: {e}")
            return False

    def _add_table_to_docx(self, doc, table_lines: List[str]):
        """添加Markdown表格到DOCX"""
        try:
            from docx.shared import Pt

            # Parse table
            rows = []
            for line in table_lines:
                if '---' in line:
                    continue
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if cells:
                    rows.append(cells)

            if not rows:
                return

            # Create table
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'

            for i, row_cells in enumerate(rows):
                row = table.rows[i]
                for j, cell_text in enumerate(row_cells):
                    if j < len(row.cells):
                        row.cells[j].text = cell_text

        except Exception as e:
            print(f"[SpecGenerator] Error adding table: {e}")

    def _generate_pdf(self, content: str, filepath: str, title: str) -> bool:
        """从Markdown内容生成PDF文档"""
        try:
            # 方案：先将Markdown转为HTML，再用Playwright生成PDF
            import markdown

            # 转换Markdown为HTML
            html_content = markdown.markdown(
                content,
                extensions=['tables', 'fenced_code', 'toc']
            )

            # 构建完整HTML
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>{title}</title>
                <style>
                    body {{
                        font-family: "Microsoft YaHei", "SimHei", sans-serif;
                        line-height: 1.8;
                        margin: 40px;
                        color: #333;
                    }}
                    h1 {{
                        color: #1a365d;
                        border-bottom: 2px solid #2b6cb0;
                        padding-bottom: 10px;
                        text-align: center;
                    }}
                    h2 {{
                        color: #2b6cb0;
                        margin-top: 30px;
                        border-left: 4px solid #2b6cb0;
                        padding-left: 10px;
                    }}
                    h3 {{
                        color: #4a5568;
                        margin-top: 20px;
                    }}
                    code {{
                        background-color: #f7fafc;
                        padding: 2px 6px;
                        border-radius: 3px;
                        font-family: "Courier New", monospace;
                        font-size: 0.9em;
                    }}
                    pre {{
                        background-color: #f7fafc;
                        padding: 15px;
                        border-radius: 5px;
                        overflow-x: auto;
                    }}
                    pre code {{
                        background-color: transparent;
                        padding: 0;
                    }}
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                        margin: 20px 0;
                    }}
                    th, td {{
                        border: 1px solid #e2e8f0;
                        padding: 12px;
                        text-align: left;
                    }}
                    th {{
                        background-color: #edf2f7;
                        font-weight: bold;
                    }}
                    blockquote {{
                        border-left: 4px solid #cbd5e0;
                        margin: 0;
                        padding-left: 20px;
                        color: #4a5568;
                    }}
                    ul, ol {{
                        margin: 10px 0;
                        padding-left: 30px;
                    }}
                    li {{
                        margin: 5px 0;
                    }}
                </style>
            </head>
            <body>
                <h1>{title}</h1>
                {html_content}
            </body>
            </html>
            """

            # 使用Playwright生成PDF
            from playwright.sync_api import sync_playwright

            with sync_playwright() as p:
                browser = p.chromium.launch()
                page = browser.new_page()
                page.set_content(full_html)
                page.pdf(
                    path=filepath,
                    format='A4',
                    margin={
                        'top': '2cm',
                        'bottom': '2cm',
                        'left': '2cm',
                        'right': '2cm'
                    },
                    print_background=True
                )
                browser.close()

            return True
        except ImportError:
            print("[SpecGenerator] Playwright not installed, falling back to markdown-pdf")
            return self._generate_pdf_fallback(content, filepath, title)
        except Exception as e:
            print(f"[SpecGenerator] Error generating PDF: {e}")
            return False

    def _generate_pdf_fallback(self, content: str, filepath: str, title: str) -> bool:
        """PDF生成降级方案 - 使用markdown-pdf"""
        try:
            # 保存markdown到临时文件
            temp_md = filepath.replace('.pdf', '_temp.md')
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(f"# {title}\n\n{content}")

            # 使用markdown-pdf（Node.js工具）
            import subprocess
            result = subprocess.run(
                ['npx', 'markdown-pdf', temp_md, '-o', filepath],
                capture_output=True,
                text=True,
                timeout=60
            )

            # 清理临时文件
            if os.path.exists(temp_md):
                os.remove(temp_md)

            return result.returncode == 0
        except Exception as e:
            print(f"[SpecGenerator] PDF fallback failed: {e}")
            return False

    def generate_specs(self, req_id: str, formats: List[str] = None) -> Tuple[bool, str, Dict]:
        """
        为特定排期需求生成规格文档

        Args:
            req_id: 需求ID
            formats: 输出格式列表 ["md", "docx", "pdf"]，默认["md"]

        Returns:
            (success, message, output_info)
        """
        if formats is None:
            formats = ["md"]

        try:
            inputs = self._resolve_generation_inputs(req_id)
            summary_success, summary_message, summary_content = self._generate_document_content(
                "概要需求",
                self._build_section_plan("summary"),
                inputs["base_context"],
                inputs["summary_template"],
                api_key=inputs["api_key"],
                provider=inputs["provider"],
                model_name=inputs["model_name"],
                base_url=inputs["base_url"] if inputs["base_url"] else None
            )
            if not summary_success:
                return False, summary_message, {}

            detail_success, detail_message, detail_content = self._generate_document_content(
                "详细需求",
                self._build_section_plan("detail"),
                inputs["base_context"],
                inputs["detail_template"],
                api_key=inputs["api_key"],
                provider=inputs["provider"],
                model_name=inputs["model_name"],
                base_url=inputs["base_url"] if inputs["base_url"] else None
            )
            if not detail_success:
                return False, detail_message, {}

            output_info = self._write_output_files(
                summary_content=summary_content,
                detail_content=detail_content,
                formats=formats,
                base_filename=inputs["base_filename"],
                safe_title=inputs["safe_title"],
                version=inputs["version"],
            )

            return True, f"Generated {len(output_info['output_files'])} files: {', '.join(output_info['output_files'])}", output_info

        except ValueError as e:
            return False, str(e), {}
        except Exception as e:
            print(f"[SpecGenerator] Error generating specs: {e}")
            return False, str(e), {}

    def get_task_status(self, task_id: str) -> Optional[Dict]:
        """获取异步任务状态"""
        task = self.tasks.get(task_id)
        if not task:
            return None
        self._summarize_task_status(task)
        return task


# Singleton instance
_spec_generator = None

def get_spec_generator(vector_store: VectorStore = None, llm_service: LLMService = None):
    """获取或创建SpecGenerator单例"""
    global _spec_generator
    if _spec_generator is None:
        if vector_store is None or llm_service is None:
            raise ValueError("vector_store and llm_service required for first initialization")
        _spec_generator = SpecGenerator(vector_store, llm_service)
    return _spec_generator
