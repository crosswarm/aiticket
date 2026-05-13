"""
报告导出服务 - Export Service
支持周报/月报的多种格式导出：Markdown、DOCX、PDF、Excel
"""

import os
import json
import uuid
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from enum import Enum
from dataclasses import dataclass, asdict
from urllib.parse import quote

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.normpath(os.path.join(BASE_DIR, "../.."))

# Export directories
EXPORT_DIR = os.path.join(PROJECT_ROOT, "conclusion/exports")
WEEKLY_REPORTS_DIR = os.path.join(PROJECT_ROOT, "conclusion/WeeklyReports")
MONTHLY_REPORTS_DIR = os.path.join(PROJECT_ROOT, "conclusion/MonthlyReports")

os.makedirs(EXPORT_DIR, exist_ok=True)


class ExportStatus(Enum):
    PENDING = "pending"
    RUNNING = "running"
    COMPLETED = "completed"
    FAILED = "failed"


@dataclass
class ExportTask:
    task_id: str
    report_type: str  # "weekly" | "monthly"
    report_id: str
    formats: List[str]
    status: ExportStatus
    progress: int
    message: str
    output_files: List[str]
    created_at: str
    updated_at: str
    error_message: Optional[str] = None
    app_base_url: Optional[str] = None


class ExportService:
    """报告导出服务"""

    def __init__(self, llm_service=None):
        self.llm_service = llm_service
        self.tasks: Dict[str, ExportTask] = {}
        self._load_task_state()

    def _load_task_state(self):
        """加载任务状态"""
        state_file = os.path.join(EXPORT_DIR, "export_tasks.json")
        if os.path.exists(state_file):
            try:
                with open(state_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    for task_id, task_data in data.items():
                        task_data['status'] = ExportStatus(task_data['status'])
                        self.tasks[task_id] = ExportTask(**task_data)
            except Exception as e:
                print(f"[ExportService] Error loading task state: {e}")

    def _save_task_state(self):
        """保存任务状态"""
        state_file = os.path.join(EXPORT_DIR, "export_tasks.json")
        try:
            data = {}
            for task_id, task in self.tasks.items():
                task_dict = asdict(task)
                task_dict['status'] = task.status.value
                data[task_id] = task_dict
            with open(state_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"[ExportService] Error saving task state: {e}")

    def start_export(self, report_type: str, report_id: str, formats: List[str], app_base_url: Optional[str] = None) -> str:
        """启动导出任务"""
        task_id = str(uuid.uuid4())[:8]

        # 验证格式
        valid_formats = ["md", "docx", "pdf", "xlsx"]
        invalid_formats = [f for f in formats if f not in valid_formats]
        if invalid_formats:
            raise ValueError(f"Invalid formats: {invalid_formats}")

        task = ExportTask(
            task_id=task_id,
            report_type=report_type,
            report_id=report_id,
            formats=formats,
            status=ExportStatus.PENDING,
            progress=0,
            message="任务已创建，等待执行",
            output_files=[],
            created_at=datetime.now().isoformat(),
            updated_at=datetime.now().isoformat(),
            app_base_url=app_base_url.rstrip("/") if app_base_url else None,
        )

        self.tasks[task_id] = task
        self._save_task_state()

        # 异步执行导出
        import threading
        thread = threading.Thread(
            target=self._execute_export,
            args=(task_id,)
        )
        thread.daemon = True
        thread.start()

        return task_id

    def _execute_export(self, task_id: str):
        """执行导出任务"""
        task = self.tasks.get(task_id)
        if not task:
            return

        try:
            task.status = ExportStatus.RUNNING
            task.message = "正在读取报告数据..."
            task.updated_at = datetime.now().isoformat()
            self._save_task_state()

            # 读取报告内容
            content, metadata = self._load_report(task.report_type, task.report_id)
            if not content:
                raise Exception("Report not found or cannot be read")

            task.progress = 20
            task.message = "正在生成导出文件..."
            task.updated_at = datetime.now().isoformat()
            self._save_task_state()

            # 生成各格式文件
            output_files = []
            base_name = f"{task.report_type}_{task.report_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

            for i, fmt in enumerate(task.formats):
                progress = 20 + (60 * (i + 1) // len(task.formats))
                task.progress = progress
                task.message = f"正在生成 {fmt.upper()} 格式..."
                task.updated_at = datetime.now().isoformat()
                self._save_task_state()

                if fmt == "md":
                    filepath = os.path.join(EXPORT_DIR, f"{base_name}.md")
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(content)
                    output_files.append(f"{base_name}.md")

                elif fmt == "docx":
                    filepath = os.path.join(EXPORT_DIR, f"{base_name}.docx")
                    if self._generate_docx(content, filepath, metadata.get('title', 'Report')):
                        output_files.append(f"{base_name}.docx")

                elif fmt == "pdf":
                    filepath = os.path.join(EXPORT_DIR, f"{base_name}.pdf")
                    if self._generate_pdf(
                        report_type=task.report_type,
                        report_id=task.report_id,
                        content=content,
                        filepath=filepath,
                        title=metadata.get('title', 'Report'),
                        app_base_url=task.app_base_url,
                    ):
                        output_files.append(f"{base_name}.pdf")

                elif fmt == "xlsx":
                    filepath = os.path.join(EXPORT_DIR, f"{base_name}.xlsx")
                    if self._generate_excel(content, filepath, metadata):
                        output_files.append(f"{base_name}.xlsx")

            task.output_files = output_files
            task.status = ExportStatus.COMPLETED
            task.progress = 100

            # 记录失败的格式
            failed_formats = [f for f in task.formats if f not in [self._get_format_from_file(of) for of in output_files]]
            if failed_formats:
                task.message = f"导出完成，成功 {len(output_files)} 个文件。失败: {', '.join(failed_formats)}"
            else:
                task.message = f"导出完成，共生成 {len(output_files)} 个文件"
            task.updated_at = datetime.now().isoformat()
            self._save_task_state()

        except Exception as e:
            print(f"[ExportService] Export failed: {e}")
            import traceback
            traceback.print_exc()
            task.status = ExportStatus.FAILED
            task.error_message = str(e)
            task.message = f"导出失败: {str(e)}"
            task.updated_at = datetime.now().isoformat()
            self._save_task_state()

    def _get_format_from_file(self, filename: str) -> str:
        """从文件名获取格式"""
        return filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    def _load_report(self, report_type: str, report_id: str) -> Tuple[str, Dict]:
        """加载报告内容"""
        metadata = {"title": "Report", "type": report_type}

        if report_type == "weekly":
            # 周报文件路径
            base_path = os.path.join(WEEKLY_REPORTS_DIR, report_id)

            # 优先处理JSON文件（新格式）
            if report_id.endswith('.json') and os.path.exists(base_path):
                with open(base_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                # 优先使用content字段（markdown格式内容）
                content = data.get('content', data.get('markdown', ''))
                metadata['title'] = os.path.basename(report_id).replace('.json', '').replace('.md', '')
                metadata['meta'] = data.get('meta', {})
                return content, metadata

            # 尝试MD文件
            if report_id.endswith('.md') and os.path.exists(base_path):
                with open(base_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                metadata['title'] = os.path.basename(report_id).replace('.md', '')
                return content, metadata

            # 尝试不带扩展名的情况
            json_path = base_path + '.json' if not base_path.endswith('.json') else base_path
            md_path = base_path + '.md' if not base_path.endswith('.md') else base_path

            if os.path.exists(json_path):
                with open(json_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                content = data.get('content', data.get('markdown', ''))
                metadata['title'] = os.path.basename(json_path).replace('.json', '')
                metadata['meta'] = data.get('meta', {})
                return content, metadata

            if os.path.exists(md_path):
                with open(md_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                metadata['title'] = os.path.basename(md_path).replace('.md', '')
                return content, metadata

        elif report_type == "monthly":
            # 月报文件路径
            base_path = os.path.join(MONTHLY_REPORTS_DIR, report_id)

            # 优先处理JSON文件
            if report_id.endswith('.json') and os.path.exists(base_path):
                with open(base_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                content = data.get('content', data.get('markdown', ''))
                metadata['title'] = os.path.basename(report_id).replace('.json', '').replace('.md', '')
                metadata['meta'] = data.get('meta', {})
                return content, metadata

            # 尝试MD文件
            if report_id.endswith('.md') and os.path.exists(base_path):
                with open(base_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                metadata['title'] = os.path.basename(report_id).replace('.md', '')
                return content, metadata

            # 尝试不带扩展名的情况
            json_path = base_path + '.json' if not base_path.endswith('.json') else base_path
            md_path = base_path + '.md' if not base_path.endswith('.md') else base_path

            if os.path.exists(json_path):
                with open(json_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                content = data.get('content', data.get('markdown', ''))
                metadata['title'] = os.path.basename(json_path).replace('.json', '')
                metadata['meta'] = data.get('meta', {})
                return content, metadata

            if os.path.exists(md_path):
                with open(md_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                metadata['title'] = os.path.basename(md_path).replace('.md', '')
                return content, metadata

        return "", metadata

    def _generate_docx(self, content: str, filepath: str, title: str) -> bool:
        """生成DOCX文档"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Add title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Parse markdown content
            lines = content.split('\n')
            current_table = []
            in_code_block = False
            code_content = []

            for line in lines:
                if line.strip().startswith('```'):
                    if in_code_block:
                        if code_content:
                            para = doc.add_paragraph()
                            run = para.add_run('\n'.join(code_content))
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                        code_content = []
                        in_code_block = False
                    else:
                        in_code_block = True
                    continue

                if in_code_block:
                    code_content.append(line)
                    continue

                # Handle headers
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                # Handle tables
                elif line.startswith('|') and '|' in line[1:]:
                    current_table.append(line)
                else:
                    if current_table:
                        self._add_table_to_docx(doc, current_table)
                        current_table = []

                    # Handle lists and paragraphs
                    if line.strip().startswith('- '):
                        doc.add_paragraph(line.strip()[2:], style='List Bullet')
                    elif re.match(r'^\d+\. ', line.strip()):
                        doc.add_paragraph(re.sub(r'^\d+\. ', '', line.strip()), style='List Number')
                    elif line.strip():
                        doc.add_paragraph(line.strip())

            if current_table:
                self._add_table_to_docx(doc, current_table)

            doc.save(filepath)
            return True
        except Exception as e:
            print(f"[ExportService] DOCX generation failed: {e}")
            return False

    def _add_table_to_docx(self, doc, table_lines: List[str]):
        """添加表格到DOCX"""
        try:
            from docx.shared import Pt

            rows = []
            for line in table_lines:
                if '---' in line:
                    continue
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if cells:
                    rows.append(cells)

            if not rows:
                return

            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'

            for i, row_cells in enumerate(rows):
                row = table.rows[i]
                for j, cell_text in enumerate(row_cells):
                    if j < len(row.cells):
                        row.cells[j].text = cell_text
        except Exception as e:
            print(f"[ExportService] Table error: {e}")

    def _build_report_print_url(self, report_type: str, report_id: str, app_base_url: str) -> str:
        encoded_report_id = quote(report_id, safe="")
        return f"{app_base_url}/report.html?print=1&type={report_type}&id={encoded_report_id}"

    def _generate_pdf_from_print_view(
        self,
        report_type: str,
        report_id: str,
        filepath: str,
        app_base_url: Optional[str],
    ) -> bool:
        if not app_base_url:
            return False

        try:
            from playwright.sync_api import sync_playwright

            target_url = self._build_report_print_url(report_type, report_id, app_base_url)

            with sync_playwright() as p:
                browser = p.chromium.launch()
                page = browser.new_page(viewport={"width": 1440, "height": 2200}, device_scale_factor=1)
                page.goto(target_url, wait_until="networkidle", timeout=60000)
                page.wait_for_function("window.__REPORT_PRINT_READY__ === true", timeout=60000)
                page.emulate_media(media="print")
                page.pdf(
                    path=filepath,
                    format="A4",
                    margin={"top": "12mm", "bottom": "12mm", "left": "12mm", "right": "12mm"},
                    print_background=True,
                    prefer_css_page_size=True,
                )
                browser.close()

            return True
        except ImportError:
            print("[ExportService] Playwright not installed")
            return False
        except Exception as e:
            print(f"[ExportService] Print-view PDF generation failed: {e}")
            return False

    def _generate_pdf_from_markdown(self, content: str, filepath: str, title: str) -> bool:
        """生成PDF文档（兼容旧导出路径）"""
        try:
            import markdown

            # Convert markdown to HTML
            html_content = markdown.markdown(
                content,
                extensions=['tables', 'fenced_code', 'toc']
            )

            # Build complete HTML with styles - optimized for A4 paper
            full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        * {{
            box-sizing: border-box;
            max-width: 100%;
        }}
        body {{
            font-family: "Microsoft YaHei", "SimHei", sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            color: #333;
            font-size: 12px;
            max-width: 190mm;
        }}
        h1 {{
            color: #1a365d;
            border-bottom: 2px solid #2b6cb0;
            padding-bottom: 8px;
            text-align: center;
            font-size: 18px;
            margin-bottom: 15px;
        }}
        h2 {{
            color: #2b6cb0;
            margin-top: 20px;
            border-left: 4px solid #2b6cb0;
            padding-left: 8px;
            font-size: 14px;
        }}
        h3 {{
            color: #4a5568;
            margin-top: 15px;
            font-size: 13px;
        }}
        code {{
            background-color: #f7fafc;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: "Courier New", monospace;
            font-size: 11px;
        }}
        pre {{
            background-color: #f7fafc;
            padding: 10px;
            border-radius: 4px;
            overflow-x: auto;
            font-size: 10px;
            white-space: pre-wrap;
            word-wrap: break-word;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 15px 0;
            font-size: 10px;
            table-layout: fixed;
        }}
        th, td {{
            border: 1px solid #e2e8f0;
            padding: 6px 8px;
            text-align: left;
            word-wrap: break-word;
            overflow-wrap: break-word;
        }}
        th {{
            background-color: #edf2f7;
            font-weight: bold;
        }}
        p {{
            margin: 8px 0;
        }}
        ul, ol {{
            margin: 8px 0;
            padding-left: 20px;
        }}
        li {{
            margin: 4px 0;
        }}
        strong {{
            color: #1a365d;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    {html_content}
</body>
</html>"""

            # Use Playwright to generate PDF
            from playwright.sync_api import sync_playwright

            with sync_playwright() as p:
                browser = p.chromium.launch()
                page = browser.new_page()
                page.set_content(full_html)
                page.pdf(
                    path=filepath,
                    format='A4',
                    margin={'top': '15mm', 'bottom': '15mm', 'left': '15mm', 'right': '15mm'},
                    print_background=True
                )
                browser.close()

            return True
        except ImportError:
            print("[ExportService] Playwright not installed")
            return False
        except Exception as e:
            print(f"[ExportService] PDF generation failed: {e}")
            return False

    def _generate_pdf(
        self,
        report_type: str,
        report_id: str,
        content: str,
        filepath: str,
        title: str,
        app_base_url: Optional[str] = None,
    ) -> bool:
        """优先使用打印视图导出 PDF，失败时回退旧 Markdown 路径"""
        if self._generate_pdf_from_print_view(report_type, report_id, filepath, app_base_url):
            return True
        return self._generate_pdf_from_markdown(content, filepath, title)

    def _generate_excel(self, content: str, filepath: str, metadata: Dict) -> bool:
        """生成Excel文档"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from openpyxl.utils import get_column_letter

            wb = Workbook()
            ws = wb.active
            ws.title = "报告摘要"

            # Parse markdown and extract data
            lines = content.split('\n')

            # Title
            ws['A1'] = metadata.get('title', '报告')
            ws['A1'].font = Font(size=16, bold=True)
            ws['A1'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A1:D1')

            current_row = 3
            current_sheet = ws

            for line in lines:
                # Headers become section titles
                if line.startswith('# '):
                    current_row += 1
                    current_sheet.cell(current_row, 1, line[2:])
                    current_sheet.cell(current_row, 1).font = Font(size=14, bold=True, color="FFFFFF")
                    current_sheet.cell(current_row, 1).fill = PatternFill(start_color="2B6CB0", end_color="2B6CB0", fill_type="solid")
                    current_row += 1

                elif line.startswith('## '):
                    current_row += 1
                    current_sheet.cell(current_row, 1, line[3:])
                    current_sheet.cell(current_row, 1).font = Font(size=12, bold=True)
                    current_row += 1

                elif line.startswith('### '):
                    current_sheet.cell(current_row, 1, line[4:])
                    current_sheet.cell(current_row, 1).font = Font(size=11, bold=True)
                    current_row += 1

                elif line.strip().startswith('- '):
                    current_sheet.cell(current_row, 2, '• ' + line.strip()[2:])
                    current_row += 1

                elif line.strip() and not line.startswith('|'):
                    current_sheet.cell(current_row, 1, line.strip())
                    current_row += 1

            # Adjust column widths
            for col in range(1, 5):
                current_sheet.column_dimensions[get_column_letter(col)].width = 20

            wb.save(filepath)
            return True
        except ImportError:
            print("[ExportService] openpyxl not installed")
            return False
        except Exception as e:
            print(f"[ExportService] Excel generation failed: {e}")
            return False

    def get_task_status(self, task_id: str) -> Optional[Dict]:
        """获取任务状态"""
        task = self.tasks.get(task_id)
        if not task:
            return None

        return {
            "task_id": task.task_id,
            "report_type": task.report_type,
            "report_id": task.report_id,
            "formats": task.formats,
            "status": task.status.value,
            "progress": task.progress,
            "message": task.message,
            "output_files": task.output_files,
            "created_at": task.created_at,
            "updated_at": task.updated_at,
            "error_message": task.error_message
        }

    def get_export_file(self, task_id: str, filename: str) -> Optional[str]:
        """获取导出文件路径"""
        filepath = os.path.join(EXPORT_DIR, filename)
        if os.path.exists(filepath):
            return filepath
        return None


# Singleton instance
_export_service = None

def get_export_service(llm_service=None) -> ExportService:
    """获取ExportService单例"""
    global _export_service
    if _export_service is None:
        _export_service = ExportService(llm_service)
    return _export_service
