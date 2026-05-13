"""附件解析器：统一处理 docx/pdf/xlsx，为 Gate 1 信息完整性检查提供文本内容。"""

from __future__ import annotations

import io
import signal
from dataclasses import dataclass, field


@dataclass
class AttachmentContent:
    text: str = ""
    tables: list[str] = field(default_factory=list)
    truncated: bool = False
    parse_error: str = ""


def _timeout_handler(signum, frame):
    raise TimeoutError("attachment parse timeout")


def _table_to_markdown(rows: list[list]) -> str:
    """二维 list 转 markdown 表格。"""
    if not rows:
        return ""

    def cell_str(c) -> str:
        if c is None:
            return ""
        return str(c).replace("\n", " ").replace("|", "\\|")

    lines: list[str] = []
    header = rows[0]
    lines.append("| " + " | ".join(cell_str(c) for c in header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows[1:]:
        # 确保列数与 header 对齐
        padded = list(row) + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(cell_str(c) for c in padded[: len(header)]) + " |")
    return "\n".join(lines)


def _parse_docx(content_bytes: bytes) -> AttachmentContent:
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(content_bytes))

    # 段落文本
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    # 表格
    tables: list[str] = []
    for tbl in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        md = _table_to_markdown(rows)
        if md:
            tables.append(md)

    truncated = False
    if len(text) > 8000:
        text = text[:8000]
        truncated = True

    return AttachmentContent(text=text, tables=tables, truncated=truncated, parse_error="")


def _parse_pdf(content_bytes: bytes) -> AttachmentContent:
    import pdfplumber

    MAX_PAGES = 50
    MAX_CHARS = 10000

    with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
        total_pages = len(pdf.pages)
        truncated = total_pages > MAX_PAGES
        pages = pdf.pages[:MAX_PAGES]

        text_parts: list[str] = []
        tables: list[str] = []

        for page in pages:
            page_text = page.extract_text() or ""
            if page_text:
                text_parts.append(page_text)

            for tbl in page.extract_tables():
                if tbl:
                    md = _table_to_markdown(tbl)
                    if md:
                        tables.append(md)

    text = "\n".join(text_parts)
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS]
        truncated = True

    return AttachmentContent(text=text, tables=tables, truncated=truncated, parse_error="")


def _parse_xlsx(content_bytes: bytes) -> AttachmentContent:
    from openpyxl import load_workbook

    MAX_ROWS = 1000

    wb = load_workbook(filename=io.BytesIO(content_bytes), data_only=True, read_only=True)
    ws = wb.worksheets[0]

    rows_data: list[list] = []
    truncated = False

    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i >= MAX_ROWS:
            truncated = True
            break
        rows_data.append(list(row))

    wb.close()

    if not rows_data:
        return AttachmentContent(text="", tables=[], truncated=False, parse_error="")

    md = _table_to_markdown(rows_data)
    return AttachmentContent(text="", tables=[md] if md else [], truncated=truncated, parse_error="")


_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif"}


def parse_attachment(url: str, filename: str, content_bytes: bytes) -> AttachmentContent:
    """
    根据文件名后缀选择解析方式，5 秒超时，超时返回空 AttachmentContent(truncated=False, parse_error='timeout')。
    """
    lower = filename.lower()

    # 图片：直接跳过，由调用方走 vision LLM
    if any(lower.endswith(ext) for ext in _IMAGE_EXTENSIONS):
        return AttachmentContent(text="", tables=[], truncated=False, parse_error="image_skip")

    # 设置 5s 超时（仅 Unix 生效）
    old_handler = signal.signal(signal.SIGALRM, _timeout_handler)
    signal.alarm(5)

    try:
        if lower.endswith(".docx"):
            result = _parse_docx(content_bytes)
        elif lower.endswith(".pdf"):
            result = _parse_pdf(content_bytes)
        elif lower.endswith(".xlsx"):
            result = _parse_xlsx(content_bytes)
        else:
            result = AttachmentContent(text="", tables=[], truncated=False, parse_error="unsupported_format")
    except TimeoutError:
        result = AttachmentContent(text="", tables=[], truncated=False, parse_error="timeout")
    except Exception as exc:  # noqa: BLE001
        result = AttachmentContent(text="", tables=[], truncated=False, parse_error=str(exc))
    finally:
        signal.alarm(0)
        signal.signal(signal.SIGALRM, old_handler)

    return result
