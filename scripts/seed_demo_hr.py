#!/usr/bin/env python3
"""
一键生成 HR 演示数据（dist 分支 Excel 数据源演示）

生成：
  data/imports/demo_hr_tickets.xlsx  — 16条HR工单（非Jira字段风格）
  KB/hr/01_attendance_policy.txt     — 考勤制度（txt格式）
  KB/hr/02_leave_application_guide.pdf  — 请假流程（pdf格式，同时生成.md用于KB索引）
  KB/hr/02_leave_application_guide.md   — 请假流程（md格式，KB直接索引）
  KB/hr/03_onboarding_checklist.docx    — 入职手册（docx格式）
  KB/hr/04_salary_components.xlsx       — 薪酬构成（xlsx格式）
  KB/hr/README.md

用法：
  python3.12 scripts/seed_demo_hr.py
  python3.12 scripts/seed_demo_hr.py --force   # 强制覆盖已有文件
"""
import argparse
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent


def _ensure_deps():
    pkgs = []
    try:
        import pandas  # noqa
    except ImportError:
        pkgs.append("pandas openpyxl")
    try:
        import docx  # noqa
    except ImportError:
        pkgs.append("python-docx")
    try:
        import reportlab  # noqa
    except ImportError:
        pkgs.append("reportlab")
    if pkgs:
        print(f"[seed] 安装依赖: {' '.join(pkgs)}")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "--quiet"] + pkgs)


def ensure_dirs():
    for d in [ROOT / "data/imports", ROOT / "KB/hr", ROOT / "samples", ROOT / "tests/demo"]:
        d.mkdir(parents=True, exist_ok=True)


# ── 1. demo_hr_tickets.xlsx ──────────────────────────────────────────────────

def _ticket_rows():
    """16 条 HR 工单数据，涵盖考勤/请假/入职/薪酬四个主题"""
    base_cols = ["工单编号", "提出者", "提出者联系方式", "所属企业",
                 "问题标题", "问题描述", "提出时间", "当前状态",
                 "处理人", "优先级", "工单类型", "业务模块"]
    rows = [
        # ── 考勤 ──────────────────────────────────────────────────────────────
        ["HR-001", "张伟", "13900139001", "上海某科技有限公司",
         "本月迟到几次才开始扣款",
         "我这个月迟到了2次，HR说还没到扣款线，想确认一下公司的迟到扣款规定是什么，从第几次迟到开始扣，每次扣多少钱。",
         "2026-05-06 09:15", "待处理", "人事-小赵", "中", "咨询", "考勤"],

        ["HR-002", "李娜", "13800138002", "北京某贸易有限公司",
         "打卡异常如何申诉",
         "上周三我在公司打卡但系统没有记录，导致考勤显示缺勤。请问考勤异常应该怎么申诉，有没有截止时间？",
         "2026-05-07 10:30", "处理中", "人事-小赵", "高", "申请", "考勤"],

        ["HR-003", "王芳", "18600001234", "广州某制造有限公司",
         "迟到30分钟和迟到10分钟有区别吗",
         "请问迟到30分钟和迟到10分钟在考勤处理上是否有区别？听说超过30分钟算旷工，想了解具体规定。",
         "2026-05-08 14:20", "已解决", "人事-小李", "低", "咨询", "考勤"],

        ["HR-004", "陈强", "17700002345", "深圳某互联网有限公司",
         "每月考勤数据什么时候统计",
         "想了解每月考勤数据的统计截止日期，以及什么时候可以看到当月的最终考勤结果。",
         "2026-05-09 08:45", "待处理", "人事-小李", "低", "咨询", "考勤"],

        # ── 请假 ──────────────────────────────────────────────────────────────
        ["HR-005", "刘洋", "13600003456", "成都某软件有限公司",
         "入职满一年有多少天年假",
         "我入职刚好满一年，想请年假，但不确定能享受多少天年假，请帮我确认一下。",
         "2026-05-10 11:00", "待处理", "人事-小赵", "中", "咨询", "请假"],

        ["HR-006", "赵敏", "13500004567", "杭州某电商有限公司",
         "请假流程是什么",
         "我需要请3天病假，不清楚具体的请假申请流程，是在钉钉上还是HR系统里申请？需要哪些材料？",
         "2026-05-11 09:30", "处理中", "人事-小李", "高", "申请", "请假"],

        ["HR-007", "周雷", "13400005678", "武汉某金融有限公司",
         "调休可以累积多少天",
         "我最近加班比较多，积累了不少调休，请问调休最多可以累积多少天，超过之后会怎么处理？",
         "2026-05-12 16:00", "已解决", "人事-小赵", "中", "咨询", "请假"],

        ["HR-008", "吴静", "18900006789", "南京某医疗有限公司",
         "病假需要提供什么证明",
         "我发烧需要请病假，请问病假需要提供哪些证明材料？超过几天需要额外手续？",
         "2026-05-13 08:00", "待处理", "人事-小李", "高", "申请", "请假"],

        # ── 入职 ──────────────────────────────────────────────────────────────
        ["HR-009", "孙浩", "13200007890", "西安某科技有限公司",
         "入职第一天需要带哪些材料",
         "我下周一正式入职，想了解报到当天需要携带什么材料，是否需要提前准备什么，以及当天会领到工牌吗？",
         "2026-05-14 10:15", "待处理", "人事-小赵", "中", "咨询", "入职"],

        ["HR-010", "马云云", "18700008901", "重庆某制造有限公司",
         "入职后多久开始缴社保",
         "我刚入职，不清楚社保是从入职当月开始缴还是下个月开始，公积金是同步的吗？",
         "2026-05-15 14:00", "处理中", "人事-小李", "中", "咨询", "入职"],

        ["HR-011", "何梅", "13100009012", "苏州某物流有限公司",
         "劳动合同什么时候签",
         "入职已经3天了，还没有签劳动合同，这正常吗？应该找谁签，什么时候截止？",
         "2026-05-16 09:45", "已解决", "人事-小赵", "高", "咨询", "入职"],

        ["HR-012", "郑伟", "18600000000", "天津某贸易有限公司",
         "试用期多长时间，转正考核怎么进行",
         "入职时HR说试用期3个月，想了解试用期内会有哪些考核，以及转正流程是什么。",
         "2026-05-17 11:30", "待处理", "人事-小李", "中", "咨询", "入职"],

        # ── 薪酬 ──────────────────────────────────────────────────────────────
        ["HR-013", "林志强", "13000001111", "厦门某互联网有限公司",
         "绩效工资是怎么计算的",
         "我的薪酬构成中有绩效工资，不清楚绩效系数是怎么评定的，绩效工资占总薪酬多少比例？",
         "2026-05-18 10:00", "待处理", "人事-小赵", "中", "咨询", "薪酬"],

        ["HR-014", "黄敏", "13100002222", "福州某科技有限公司",
         "周末加班费怎么算",
         "最近公司安排周末加班，想了解加班费的计算标准，周末加班和法定节假日加班有什么区别？",
         "2026-05-19 09:00", "处理中", "人事-小李", "高", "咨询", "薪酬"],

        ["HR-015", "曹磊", "13200003333", "合肥某制造有限公司",
         "个税怎么扣，有什么专项扣除",
         "想了解工资中的个人所得税是怎么计算和扣缴的，有哪些专项附加扣除可以申报？",
         "2026-05-19 15:30", "待处理", "人事-小赵", "低", "咨询", "薪酬"],

        ["HR-016", "田娟", "13300004444", "无锡某电商有限公司",
         "年终奖什么时候发，怎么计算",
         "听说公司有年终奖，请问年终奖一般什么时候发放，计算标准是什么，跟绩效挂钩吗？",
         "2026-05-20 08:30", "待处理", "人事-小李", "中", "咨询", "薪酬"],
    ]
    return base_cols, rows


def generate_tickets_xlsx(force=False):
    import pandas as pd
    out = ROOT / "data/imports/demo_hr_tickets.xlsx"
    if out.exists() and not force:
        print(f"[seed] 跳过（已存在）: {out.relative_to(ROOT)}")
        return
    cols, rows = _ticket_rows()
    df = pd.DataFrame(rows, columns=cols)
    # 强制所有列为字符串（避免日期列被解析）
    df = df.astype(str)
    df.to_excel(out, index=False, engine="openpyxl")
    print(f"[seed] ✅ 生成: {out.relative_to(ROOT)} ({len(rows)} 条工单)")


# ── 2. KB/hr/01_attendance_policy.txt ───────────────────────────────────────

ATTENDANCE_CONTENT = """\
# HR管理系统 - 考勤管理制度

## 一、工作时间

正常工作时间：每天 9:00 - 18:00，午休 12:00 - 13:30。
弹性签到时间：9:00 至 10:00 之间签到均视为正常到岗。
每周工作5天，周六日为法定休息日。

## 二、迟到处理规定

1. 每月迟到 3 次以内（含3次）不做扣款处理。
2. 超过 3 次后，每次迟到扣款 50 元。
3. 迟到 30 分钟以上计旷工半天；旷工处理标准：日薪 / 2。
4. 全月旷工超过 3 天者，视情节轻重给予纪律处分。

## 三、早退处理

早退处理标准与迟到相同，累计计入当月迟到/早退次数。

## 四、考勤异常申诉

1. 考勤异常须在异常发生后 **24小时内** 提交申诉，超期不予受理。
2. 申诉渠道：通过HR系统「考勤管理」-「异常申诉」模块提交，须附打卡截图或证明材料。
3. 申诉审批周期：1-3 个工作日，结果通过邮件和HR系统通知。
4. 申诉截止日：每月 **25日** 为当月考勤截止日，25日后不再受理当月异常申诉。

## 五、月度考勤统计

- 每月 25 日为考勤数据统计截止日。
- HR 部门在每月 28 日前完成考勤数据核对，并发送至财务部门。
- 员工可在HR系统查看当月考勤汇总，如有异议请在截止日前提出。

## 六、加班登记

1. 工作日 18:00 后加班须提前申请，经直属上级审批。
2. 周末加班须提前 1 天提交加班申请。
3. 加班可选择调休（优先）或加班费（见薪酬制度）。
"""


def generate_kb_txt(force=False):
    out = ROOT / "KB/hr/01_attendance_policy.txt"
    if out.exists() and not force:
        print(f"[seed] 跳过（已存在）: {out.relative_to(ROOT)}")
        return
    out.write_text(ATTENDANCE_CONTENT, encoding="utf-8")
    print(f"[seed] ✅ 生成: {out.relative_to(ROOT)}")


# ── 3. KB/hr/02_leave_application_guide.pdf + .md ───────────────────────────

LEAVE_CONTENT = """\
# HR管理系统 - 请假申请指南

## 一、年假政策

根据《职工带薪年休假条例》及公司规定，年假天数按工龄计算：

| 工龄 | 年假天数 |
|------|--------|
| 不满1年 | 5天 |
| 满1年不满10年 | 10天 |
| 满10年以上 | 15天 |

注：工龄以社保缴纳年限为准，跨公司工龄可累积。

## 二、病假规定

1. 病假须提供正规医院（二级及以上）的诊断证明。
2. 连续病假超过 3 天需向 HR 备案并提交医院证明原件。
3. 连续病假超过 30 天，须提交复工证明方可回岗。
4. 病假期间工资按基本工资的 80% 发放。

## 三、调休管理

1. 加班调休可累积保留，最多累积 **90天**。
2. 累积调休超过 90 天的部分自动作废，不计为加班费。
3. 调休须提前 1 天提交申请，经直属上级审批。
4. 年底（12月31日）未使用的调休不予结转。

## 四、请假申请流程（重要步骤）

步骤1：登录钉钉 HR 系统，进入【假勤管理】-【请假申请】。
步骤2：选择请假类型（年假/病假/调休/事假等）和起止时间，填写请假事由。
步骤3：系统自动推送至直属上级审批（须在 1 个工作日内完成审批）。
步骤4：上级审批通过后，HR 备案（仅病假、婚假、产假需要额外备案）。
步骤5：请假结束后，如涉及调休请及时在系统中标注，确保考勤数据准确。

## 五、特殊假种说明

| 假种 | 天数 | 所需证明 |
|------|------|--------|
| 婚假 | 3天 | 结婚证 |
| 产假 | 158天（法定98天+生育奖励假60天） | 出生医学证明 |
| 陪产假 | 15天 | 出生医学证明 |
| 丧假 | 3天（直系亲属） | 死亡证明 |

## 六、注意事项

- 未经审批擅自离岗视为旷工处理。
- 超出年假剩余天数的，超出部分按事假处理（扣除相应日薪）。
- 请假记录可在钉钉 HR 系统「假勤查询」模块查询。
"""


def generate_kb_pdf(force=False):
    out_pdf = ROOT / "KB/hr/02_leave_application_guide.pdf"
    out_md = ROOT / "KB/hr/02_leave_application_guide.md"

    # 始终生成 .md 供 KB 索引
    if not out_md.exists() or force:
        out_md.write_text(LEAVE_CONTENT, encoding="utf-8")
        print(f"[seed] ✅ 生成: {out_md.relative_to(ROOT)}")

    if out_pdf.exists() and not force:
        print(f"[seed] 跳过（已存在）: {out_pdf.relative_to(ROOT)}")
        return

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm
        import reportlab.rl_config as rl_cfg

        # 尝试注册中文字体（跨平台）
        chinese_font = None
        font_candidates = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/Supplemental/Songti.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        ]
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            for fp in font_candidates:
                if Path(fp).exists():
                    pdfmetrics.registerFont(TTFont("ChineseFont", fp))
                    chinese_font = "ChineseFont"
                    break
        except Exception:
            pass

        doc = SimpleDocTemplate(str(out_pdf), pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        story = []

        font_name = chinese_font or "Helvetica"
        from reportlab.lib.styles import ParagraphStyle
        body_style = ParagraphStyle("body", parent=styles["Normal"],
                                    fontName=font_name, fontSize=10,
                                    leading=16, spaceAfter=6)
        title_style = ParagraphStyle("title", parent=styles["Heading1"],
                                     fontName=font_name, fontSize=14)

        for line in LEAVE_CONTENT.splitlines():
            if line.startswith("# "):
                story.append(Paragraph(line[2:], title_style))
            elif line.strip():
                # Escape HTML chars for reportlab
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, body_style))
            else:
                story.append(Spacer(1, 6))

        doc.build(story)
        print(f"[seed] ✅ 生成: {out_pdf.relative_to(ROOT)}")
    except Exception as e:
        print(f"[seed] ⚠️ PDF 生成失败（{e}），仅使用 .md 版本用于 KB 索引")


# ── 4. KB/hr/03_onboarding_checklist.docx ───────────────────────────────────

ONBOARDING_CONTENT_SECTIONS = [
    ("HR管理系统 - 新员工入职手册", None),
    ("一、入职前准备", [
        "收到录用通知后，登录 HR 系统完善个人信息（含银行账户、紧急联系人、学历信息等）。",
        "准备以下材料原件：身份证、毕业证、离职证明（如有）、近期一寸照片 2 张。",
        "提前了解公司位置和报到时间，确认入职当天的具体安排。",
    ]),
    ("二、报到当天（Day 1）", [
        "报到时间：上午 9:00 前至前台签到，迟到须提前通知 HR。",
        "领取物品：",
        "  - 工牌（含门禁权限，当天生效）",
        "  - 工作笔记本电脑（IT 当天完成初始化配置）",
        "  - 办公文具一套",
        "  - 公司手册及员工守则（纸质版）",
        "参加新员工入职培训（人力资源部组织，上午 10:00 开始）。",
    ]),
    ("三、入职一周内必须完成", [
        "社保登记：HR 在入职 3 个工作日内完成社保参保手续，次月起生效。",
        "公积金账户：在职当月满 15 天的，当月开始缴公积金；不满 15 天的次月起缴。",
        "签订劳动合同：入职 1 周内完成劳动合同签署（电子版，通过 HR 系统完成）。",
        "账号开通：IT 在入职 1 个工作日内完成工作邮箱、内网账号、飞书账号开通。",
        "工牌激活：在 HR 系统绑定工号，完成照片更新。",
    ]),
    ("四、试用期安排", [
        "试用期：90 天，从入职当日起计算（法定最长不超过 6 个月）。",
        "期间考核：第 30 天、第 60 天各进行一次绩效面谈，结果记入转正评估。",
        "转正评估：第 85 天前由直属上级提交转正申请，HR 审核后通知员工。",
        "试用期工资：通常为正式工资的 80%（具体以 Offer 为准）。",
    ]),
    ("五、导师（Mentor）计划", [
        "每位新员工配备 1 名 Mentor（入职前由 HR 协调配对，同部门资深员工担任）。",
        "Mentor 陪同时间：前 30 天每周至少 1 次面谈（≥30 分钟）。",
        "Mentor 职责：业务导向、系统培训、文化融入、答疑解惑。",
        "新员工培训期间提出的问题，优先由 Mentor 解答，其次联系 HR。",
    ]),
]


def generate_kb_docx(force=False):
    out = ROOT / "KB/hr/03_onboarding_checklist.docx"
    if out.exists() and not force:
        print(f"[seed] 跳过（已存在）: {out.relative_to(ROOT)}")
        return
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        for title, items in ONBOARDING_CONTENT_SECTIONS:
            if items is None:
                h = doc.add_heading(title, level=1)
            else:
                doc.add_heading(title, level=2)
                for item in items:
                    p = doc.add_paragraph(item)
                    p.style.font.size = Pt(11)
        doc.save(str(out))
        print(f"[seed] ✅ 生成: {out.relative_to(ROOT)}")
    except Exception as e:
        print(f"[seed] ❌ docx 生成失败: {e}")
        raise


# ── 5. KB/hr/04_salary_components.xlsx ──────────────────────────────────────

SALARY_ROWS = [
    ["薪酬项", "计算口径", "发放频率", "备注"],
    ["基本工资", "月薪 × 70%", "每月", "固定部分，不与绩效挂钩"],
    ["绩效工资", "月薪 × 30% × 绩效系数（0.5-1.5）", "每月", "绩效系数由直属上级与 HR 共同评定"],
    ["加班费", "时薪 × 加班时长 × 倍数", "每月", "平日 1.5 倍 / 周末 2.0 倍 / 法定节假日 3.0 倍"],
    ["年终奖", "平均月薪 × N 个月", "每年", "N 由公司根据年度经营情况决定，一般 1-3 个月"],
    ["餐补", "固定 200 元/月", "每月", "在职当月全额发放，离职当月按出勤比例发放"],
    ["交通补贴", "固定 300 元/月", "每月", "出差期间暂停发放，报销出差交通费"],
    ["个税", "按累计预扣法代扣", "每月代扣", "遵照国家税务总局最新规定，支持专项附加扣除"],
    ["养老保险", "缴费基数 × 8%（个人）", "每月", "公司缴纳 16%，个人缴纳 8%"],
    ["医疗保险", "缴费基数 × 2%（个人）", "每月", "公司缴纳 10%，个人缴纳 2%（+3元大病保险）"],
    ["失业保险", "缴费基数 × 0.5%（个人）", "每月", "公司缴纳 0.5%，个人缴纳 0.5%"],
    ["住房公积金", "缴存基数 × 12%（个人）", "每月", "公司 12%，个人 12%，共计缴存基数的 24%"],
]


def generate_kb_xlsx(force=False):
    import pandas as pd
    out = ROOT / "KB/hr/04_salary_components.xlsx"
    if out.exists() and not force:
        print(f"[seed] 跳过（已存在）: {out.relative_to(ROOT)}")
        return
    header, *data = SALARY_ROWS
    df = pd.DataFrame(data, columns=header)
    writer = pd.ExcelWriter(str(out), engine="openpyxl")
    df.to_excel(writer, index=False, sheet_name="薪酬构成")
    ws = writer.sheets["薪酬构成"]
    # 简单列宽
    ws.column_dimensions["A"].width = 16
    ws.column_dimensions["B"].width = 36
    ws.column_dimensions["C"].width = 12
    ws.column_dimensions["D"].width = 40
    writer.close()
    print(f"[seed] ✅ 生成: {out.relative_to(ROOT)}")

    # Also write a .md companion so KB can index the salary content directly
    md_out = ROOT / "KB/hr/04_salary_components.md"
    if not md_out.exists() or force:
        md_lines = ["# 薪酬结构说明\n",
                    "本文档说明公司薪酬各组成部分的计算方式与发放频率。\n",
                    "## 薪酬组成\n",
                    "| 薪酬项 | 计算口径 | 发放频率 | 备注 |\n",
                    "|--------|----------|----------|------|\n"]
        for row in data:
            md_lines.append(f"| {' | '.join(str(c) for c in row)} |\n")
        md_lines += ["\n## 绩效系数说明\n",
                     "- 优秀（S）：系数 1.5；良好（A）：1.2；合格（B）：1.0；待改进（C）：0.7；不合格（D）：0.5\n",
                     "\n## 加班费计算\n",
                     "时薪 = 月薪 ÷ 21.75 ÷ 8；平日 1.5 倍 / 周末 2.0 倍 / 法定节假日 3.0 倍\n",
                     "\n## 个税专项扣除\n",
                     "子女教育每月 1000 元 / 住房贷款利息每月 1000 元 / 赡养老人最高 2000 元\n",
                     "\n## 发放时间\n",
                     "每月 15 日前发放上月薪酬（遇节假日顺延）。\n"]
        md_out.write_text("".join(md_lines), encoding="utf-8")
        print(f"[seed] ✅ 生成: {md_out.relative_to(ROOT)}")


# ── 6. KB/hr/README.md ──────────────────────────────────────────────────────

KB_README = """\
# HR 管理系统知识库

本目录存放人力资源管理系统相关政策和操作指南，供 AI 智能回复检索使用。

## 文档清单

| 文件 | 格式 | 内容 |
|------|------|------|
| 01_attendance_policy.txt | TXT | 考勤管理制度（迟到/早退/旷工/申诉规则）|
| 02_leave_application_guide.pdf | PDF | 请假申请流程（年假/病假/调休/各类假种）|
| 02_leave_application_guide.md  | MD  | 请假指南 Markdown 版（KB 直接索引）|
| 03_onboarding_checklist.docx | DOCX | 新员工入职手册（报到/社保/试用期/导师）|
| 04_salary_components.xlsx | XLSX | 薪酬构成说明（各项计算口径和发放频率）|

## 使用方式

将此目录的内容用于 KB 知识库索引时，运行：

```bash
curl -X POST http://localhost:18000/api/kb/sync
```

或使用导入脚本：

```bash
python3 scripts/import_kb.py
```

## 演示数据配套

与 `data/imports/demo_hr_tickets.xlsx`（16 条 HR 工单）配合使用，
每条工单的智能回复应能检索到本目录对应知识文档。
"""


def generate_kb_readme(force=False):
    out = ROOT / "KB/hr/README.md"
    if out.exists() and not force:
        print(f"[seed] 跳过（已存在）: {out.relative_to(ROOT)}")
        return
    out.write_text(KB_README, encoding="utf-8")
    print(f"[seed] ✅ 生成: {out.relative_to(ROOT)}")


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="生成 HR Demo 演示数据")
    parser.add_argument("--force", action="store_true", help="强制覆盖已有文件")
    args = parser.parse_args()

    _ensure_deps()
    ensure_dirs()

    generate_tickets_xlsx(args.force)
    generate_kb_txt(args.force)
    generate_kb_pdf(args.force)
    generate_kb_docx(args.force)
    generate_kb_xlsx(args.force)
    generate_kb_readme(args.force)

    print("\n[seed] ✅ 全部完成！后续步骤：")
    print("  1. 合并 demo 配置：cp samples/deployment.demo-hr.yaml config/deployment.yaml")
    print("     （保留你已有的 llm 配置段）")
    print("  2. 启动 backend 后触发 KB 索引：")
    print("     curl -X POST http://localhost:18000/api/kb/sync")
    print("  3. 验证：python scripts/validate_demo_e2e.py")


if __name__ == "__main__":
    main()
